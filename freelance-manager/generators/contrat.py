"""
Générateur de contrats Word/PDF pour freelance.
3 types : Régie, Forfait, Mission courte.
"""

from pathlib import Path
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT

from database import Client, Contrat


class ContratGenerator:
    """Génère des contrats Word professionnels."""

    def __init__(self, config: dict):
        self.config = config
        self.freelance = config.get('freelance', {})
        self.output_dir = Path(config.get('paths', {}).get('output', 'output'))
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate(self, contrat: Contrat, client: Client,
                 description: str = "", duree_mois: int = 3,
                 lieu_mission: str = "", objet: str = "") -> str:
        """
        Génère le contrat Word.

        Args:
            contrat: Objet Contrat
            client: Objet Client
            description: Description de la mission
            duree_mois: Durée en mois
            lieu_mission: Lieu d'exécution
            objet: Objet du contrat

        Returns:
            Chemin du fichier Word généré.
        """
        if contrat.type_contrat == 'regie':
            return self._generate_regie(contrat, client, description, duree_mois, lieu_mission, objet)
        elif contrat.type_contrat == 'forfait':
            return self._generate_forfait(contrat, client, description, lieu_mission, objet)
        elif contrat.type_contrat == 'mission':
            return self._generate_mission(contrat, client, description, lieu_mission, objet)
        else:
            raise ValueError(f"Type de contrat inconnu: {contrat.type_contrat}")

    def _create_document(self) -> Document:
        """Crée un nouveau document avec styles de base."""
        doc = Document()

        # Marges
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        return doc

    def _add_header(self, doc: Document, title: str, numero: str):
        """Ajoute l'en-tête du contrat."""
        # Titre
        heading = doc.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Numéro
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Contrat N° {numero}")
        run.font.size = Pt(12)
        run.font.italic = True

        doc.add_paragraph()

    def _add_parties(self, doc: Document, client: Client):
        """Ajoute la section des parties."""
        doc.add_heading("ENTRE LES SOUSSIGNÉS", level=1)

        # Le prestataire
        p = doc.add_paragraph()
        p.add_run("Le Prestataire :").bold = True
        doc.add_paragraph(f"""
{self.freelance.get('nom', 'Nom Prénom')}
{self.freelance.get('statut', 'Micro-entrepreneur')}
SIRET : {self.freelance.get('siret', 'N/A')}
{self.freelance.get('adresse', '')}
{self.freelance.get('code_postal', '')} {self.freelance.get('ville', '')}
Email : {self.freelance.get('email', '')}
Tél : {self.freelance.get('telephone', '')}
""".strip())

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Ci-après dénommé « le Prestataire »").italic = True

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("ET")

        doc.add_paragraph()

        # Le client
        p = doc.add_paragraph()
        p.add_run("Le Client :").bold = True
        client_info = f"{client.nom}"
        if client.siret:
            client_info += f"\nSIRET : {client.siret}"
        if client.adresse:
            client_info += f"\n{client.adresse}"
        if client.code_postal or client.ville:
            client_info += f"\n{client.code_postal} {client.ville}"
        if client.contact_nom:
            client_info += f"\nReprésenté par : {client.contact_nom}"
        if client.email:
            client_info += f"\nEmail : {client.email}"

        doc.add_paragraph(client_info)

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Ci-après dénommé « le Client »").italic = True

        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run("Ensemble dénommés « les Parties »").italic = True

    def _add_signature_block(self, doc: Document):
        """Ajoute le bloc de signature."""
        doc.add_heading("SIGNATURES", level=1)

        p = doc.add_paragraph()
        p.add_run("Fait en deux exemplaires originaux,")

        p = doc.add_paragraph()
        p.add_run(f"À _________________, le _________________")

        doc.add_paragraph()
        doc.add_paragraph()

        # Tableau signatures
        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # En-têtes
        cell_prestataire = table.cell(0, 0)
        cell_prestataire.text = "Le Prestataire"
        cell_prestataire.paragraphs[0].runs[0].bold = True

        cell_client = table.cell(0, 1)
        cell_client.text = "Le Client"
        cell_client.paragraphs[0].runs[0].bold = True

        # Mention
        table.cell(1, 0).text = "(Signature précédée de la mention\n« Lu et approuvé »)"
        table.cell(1, 1).text = "(Signature précédée de la mention\n« Lu et approuvé »)"

        # Espace pour signature
        table.cell(2, 0).text = "\n\n\n"
        table.cell(2, 1).text = "\n\n\n"

        # Noms
        table.cell(3, 0).text = f"{self.freelance.get('nom', '')}"
        table.cell(3, 1).text = f"{client.nom if hasattr(client, 'nom') else ''}"

    def _generate_regie(self, contrat: Contrat, client: Client,
                        description: str, duree_mois: int,
                        lieu_mission: str, objet: str) -> str:
        """Génère un contrat de prestation en régie."""
        doc = self._create_document()

        self._add_header(doc, "CONTRAT DE PRESTATION DE SERVICES EN RÉGIE", contrat.numero)
        self._add_parties(doc, client)

        # PRÉAMBULE
        doc.add_heading("PRÉAMBULE", level=1)
        doc.add_paragraph(f"""
Le Client souhaite bénéficier des compétences du Prestataire pour la réalisation de prestations de services informatiques en mode régie.

Le Prestataire déclare disposer des compétences et de l'expérience nécessaires pour répondre aux besoins du Client.

Les Parties ont convenu ce qui suit :
""".strip())

        # ARTICLE 1 - OBJET
        doc.add_heading("ARTICLE 1 - OBJET DU CONTRAT", level=1)
        doc.add_paragraph(f"""
Le présent contrat a pour objet de définir les conditions dans lesquelles le Prestataire s'engage à fournir au Client des prestations de services en régie.

{f"Description de la mission : {objet}" if objet else "La nature exacte des prestations sera définie dans les bons de commande annexés au présent contrat."}

{f"Détail : {description}" if description else ""}
""".strip())

        # ARTICLE 2 - DURÉE
        date_debut = contrat.date_debut or date.today()
        date_fin = contrat.date_fin or (date_debut + timedelta(days=duree_mois * 30))

        doc.add_heading("ARTICLE 2 - DURÉE", level=1)
        doc.add_paragraph(f"""
Le présent contrat est conclu pour une durée déterminée.

Date de début : {date_debut.strftime('%d/%m/%Y')}
Date de fin prévisionnelle : {date_fin.strftime('%d/%m/%Y')}

Le contrat pourra être renouvelé par accord écrit des Parties.

Chaque Partie pourra mettre fin au contrat avec un préavis de 15 jours ouvrés, notifié par lettre recommandée avec accusé de réception ou par email avec accusé de réception.
""".strip())

        # ARTICLE 3 - LIEU D'EXÉCUTION
        doc.add_heading("ARTICLE 3 - LIEU D'EXÉCUTION", level=1)
        doc.add_paragraph(f"""
Les prestations seront réalisées :
{f"- {lieu_mission}" if lieu_mission else "- Dans les locaux du Client et/ou en télétravail selon les besoins de la mission"}

Le Prestataire s'engage à respecter le règlement intérieur du Client lorsqu'il intervient dans ses locaux.
""".strip())

        # ARTICLE 4 - RÉMUNÉRATION
        doc.add_heading("ARTICLE 4 - RÉMUNÉRATION", level=1)
        doc.add_paragraph(f"""
La rémunération du Prestataire est fixée sur la base d'un Taux Journalier Moyen (TJM) de :

    {contrat.tjm:,.2f} € HT / jour

{"TVA non applicable, article 293 B du CGI." if not self.freelance.get('tva_applicable', False) else "TVA au taux légal en vigueur."}

La facturation sera établie mensuellement sur la base des jours effectivement travaillés, validés par le Client.

Un compte-rendu d'activité (CRA) sera établi en fin de mois et devra être validé par le Client avant facturation.
""".strip())

        # ARTICLE 5 - CONDITIONS DE PAIEMENT
        doc.add_heading("ARTICLE 5 - CONDITIONS DE PAIEMENT", level=1)
        doc.add_paragraph("""
Les factures sont payables à 30 jours fin de mois à réception.

Le paiement sera effectué par virement bancaire aux coordonnées indiquées sur la facture.

En cas de retard de paiement, des pénalités de retard au taux légal en vigueur seront appliquées de plein droit, ainsi qu'une indemnité forfaitaire pour frais de recouvrement de 40 euros (art. L.441-10 du Code de commerce).
""".strip())

        # ARTICLE 6 - NON-EXCLUSIVITÉ
        doc.add_heading("ARTICLE 6 - NON-EXCLUSIVITÉ", level=1)
        doc.add_paragraph("""
Le Prestataire conserve son entière indépendance dans l'organisation et l'exécution de ses prestations.

Le présent contrat n'est pas exclusif. Le Prestataire reste libre d'exercer son activité pour d'autres clients, sous réserve du respect de ses engagements au titre du présent contrat et des obligations de confidentialité.

Le Prestataire s'engage toutefois à informer le Client de toute situation de conflit d'intérêts potentiel.
""".strip())

        # ARTICLE 7 - CONFIDENTIALITÉ
        doc.add_heading("ARTICLE 7 - CONFIDENTIALITÉ", level=1)
        doc.add_paragraph("""
Le Prestataire s'engage à :
- Garder strictement confidentielles toutes les informations portées à sa connaissance dans le cadre de la mission ;
- Ne pas divulguer ces informations à des tiers sans l'accord préalable écrit du Client ;
- N'utiliser ces informations que dans le cadre de l'exécution du présent contrat.

Cette obligation de confidentialité restera en vigueur pendant 2 ans après la fin du contrat.

Ne sont pas concernées par cette obligation les informations :
- Qui sont ou deviennent publiques sans faute du Prestataire ;
- Qui étaient déjà en possession du Prestataire avant leur divulgation ;
- Qui sont obtenues légitimement d'un tiers autorisé à les divulguer.
""".strip())

        # ARTICLE 8 - PROPRIÉTÉ INTELLECTUELLE
        doc.add_heading("ARTICLE 8 - PROPRIÉTÉ INTELLECTUELLE", level=1)
        doc.add_paragraph("""
Tous les travaux, documents, développements et créations réalisés par le Prestataire dans le cadre du présent contrat seront la propriété exclusive du Client, dès leur paiement intégral.

Le Prestataire cède au Client, de manière exclusive, l'ensemble des droits patrimoniaux d'auteur sur les livrables, pour toute la durée légale de protection et pour le monde entier.

Le Prestataire garantit que les livrables ne portent pas atteinte aux droits de tiers et qu'ils sont originaux.

Le Prestataire conserve le droit de mentionner sa participation au projet à titre de référence commerciale, sauf opposition écrite du Client.
""".strip())

        # ARTICLE 9 - RESPONSABILITÉ
        doc.add_heading("ARTICLE 9 - RESPONSABILITÉ", level=1)
        doc.add_paragraph("""
Le Prestataire s'engage à exécuter ses prestations avec tout le soin et la diligence requis, conformément aux règles de l'art.

La responsabilité du Prestataire est une obligation de moyens.

En aucun cas, le Prestataire ne pourra être tenu responsable :
- Des dommages indirects subis par le Client ;
- Des pertes de données, de chiffre d'affaires ou de bénéfices ;
- De tout dommage résultant d'une utilisation non conforme des livrables.

La responsabilité du Prestataire est limitée au montant des sommes perçues au titre du présent contrat sur les 12 derniers mois.
""".strip())

        # ARTICLE 10 - INDÉPENDANCE
        doc.add_heading("ARTICLE 10 - INDÉPENDANCE DES PARTIES", level=1)
        doc.add_paragraph("""
Les Parties reconnaissent expressément que le présent contrat ne crée entre elles aucun lien de subordination.

Le Prestataire exerce son activité de manière indépendante et assume seul la gestion de son entreprise.

Le Prestataire déclare être régulièrement immatriculé et à jour de ses obligations sociales et fiscales.
""".strip())

        # ARTICLE 11 - RÉSILIATION
        doc.add_heading("ARTICLE 11 - RÉSILIATION", level=1)
        doc.add_paragraph("""
En cas de manquement grave de l'une des Parties à ses obligations contractuelles, l'autre Partie pourra résilier le contrat de plein droit, 15 jours après mise en demeure restée sans effet.

En cas de résiliation anticipée :
- Le Prestataire sera rémunéré pour les prestations déjà réalisées ;
- Le Prestataire remettra au Client tous les travaux en cours.
""".strip())

        # ARTICLE 12 - LOI APPLICABLE
        doc.add_heading("ARTICLE 12 - LOI APPLICABLE ET JURIDICTION", level=1)
        doc.add_paragraph(f"""
Le présent contrat est soumis au droit français.

En cas de litige, les Parties s'engagent à rechercher une solution amiable.

À défaut d'accord, le litige sera soumis aux tribunaux compétents de {self.freelance.get('ville', 'Strasbourg')}.
""".strip())

        # SIGNATURES
        self._add_signature_block(doc)

        # Sauvegarde
        filename = f"{contrat.numero}_regie.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))

        return str(filepath)

    def _generate_forfait(self, contrat: Contrat, client: Client,
                          description: str, lieu_mission: str, objet: str) -> str:
        """Génère un contrat au forfait."""
        doc = self._create_document()

        self._add_header(doc, "CONTRAT DE PRESTATION AU FORFAIT", contrat.numero)
        self._add_parties(doc, client)

        # PRÉAMBULE
        doc.add_heading("PRÉAMBULE", level=1)
        doc.add_paragraph("""
Le Client souhaite confier au Prestataire la réalisation d'une mission définie, avec un engagement sur le résultat et un prix forfaitaire.

Le Prestataire déclare avoir pris connaissance des besoins du Client et s'engage à livrer les prestations décrites ci-après.
""".strip())

        # ARTICLE 1 - OBJET
        doc.add_heading("ARTICLE 1 - OBJET ET LIVRABLES", level=1)
        doc.add_paragraph(f"""
Le présent contrat a pour objet la réalisation par le Prestataire des prestations suivantes :

{objet if objet else "À définir dans le cahier des charges annexé."}

{f"Description détaillée : {description}" if description else ""}

Les livrables attendus sont les suivants :
- [À compléter]
- [À compléter]
- [À compléter]

Le Prestataire s'engage sur un résultat conforme aux spécifications définies.
""".strip())

        # ARTICLE 2 - PRIX FORFAITAIRE
        doc.add_heading("ARTICLE 2 - PRIX FORFAITAIRE", level=1)
        montant = contrat.montant_forfait or (contrat.tjm * (contrat.duree_jours or 10))
        doc.add_paragraph(f"""
Le prix forfaitaire de la prestation est fixé à :

    {montant:,.2f} € HT

{"TVA non applicable, article 293 B du CGI." if not self.freelance.get('tva_applicable', False) else "TVA au taux légal en vigueur."}

Ce prix est ferme et définitif. Il couvre l'ensemble des prestations décrites à l'article 1.

Toute prestation supplémentaire fera l'objet d'un avenant.
""".strip())

        # ARTICLE 3 - ÉCHÉANCIER DE PAIEMENT
        doc.add_heading("ARTICLE 3 - ÉCHÉANCIER DE PAIEMENT", level=1)
        acompte = montant * 0.3
        solde = montant - acompte
        doc.add_paragraph(f"""
Le paiement s'effectuera selon l'échéancier suivant :

1. Acompte à la signature : {acompte:,.2f} € HT (30%)
2. Solde à la livraison : {solde:,.2f} € HT (70%)

Les factures sont payables à 30 jours à réception.
""".strip())

        # ARTICLE 4 - DÉLAIS ET JALONS
        date_debut = contrat.date_debut or date.today()
        date_fin = contrat.date_fin or (date_debut + timedelta(days=(contrat.duree_jours or 20)))

        doc.add_heading("ARTICLE 4 - DÉLAIS ET JALONS", level=1)
        doc.add_paragraph(f"""
Date de début : {date_debut.strftime('%d/%m/%Y')}
Date de livraison prévue : {date_fin.strftime('%d/%m/%Y')}

Jalons intermédiaires :
- Jalon 1 : [À définir] - Date : [À définir]
- Jalon 2 : [À définir] - Date : [À définir]
- Livraison finale : {date_fin.strftime('%d/%m/%Y')}

Le Client s'engage à valider chaque jalon dans un délai de 5 jours ouvrés.
""".strip())

        # ARTICLE 5 - RECETTE ET VALIDATION
        doc.add_heading("ARTICLE 5 - RECETTE ET VALIDATION", level=1)
        doc.add_paragraph("""
À la livraison de chaque jalon, le Client dispose d'un délai de 5 jours ouvrés pour :
- Valider la livraison, ou
- Formuler ses réserves par écrit

En l'absence de réponse dans ce délai, la livraison sera réputée acceptée.

En cas de réserves, le Prestataire dispose de 5 jours ouvrés pour effectuer les corrections. Une nouvelle recette sera alors organisée.
""".strip())

        # ARTICLE 6 - PÉNALITÉS DE RETARD
        doc.add_heading("ARTICLE 6 - PÉNALITÉS DE RETARD", level=1)
        doc.add_paragraph("""
En cas de retard imputable au Prestataire dans la livraison des livrables, des pénalités de retard pourront être appliquées à hauteur de 1% du montant forfaitaire par semaine de retard, plafonnées à 10% du montant total.

Ces pénalités ne seront pas applicables si le retard est imputable au Client (retard de validation, changement de spécifications, etc.).
""".strip())

        # Articles standards (confidentialité, PI, etc.)
        self._add_standard_articles(doc)

        # SIGNATURES
        self._add_signature_block(doc)

        # Sauvegarde
        filename = f"{contrat.numero}_forfait.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))

        return str(filepath)

    def _generate_mission(self, contrat: Contrat, client: Client,
                          description: str, lieu_mission: str, objet: str) -> str:
        """Génère un contrat de mission courte (1-5 jours)."""
        doc = self._create_document()

        self._add_header(doc, "CONTRAT DE MISSION", contrat.numero)
        self._add_parties(doc, client)

        # OBJET
        doc.add_heading("ARTICLE 1 - MISSION", level=1)
        date_debut = contrat.date_debut or date.today()
        jours = contrat.duree_jours or 1

        doc.add_paragraph(f"""
Le Prestataire s'engage à réaliser pour le Client la mission suivante :

{objet if objet else "Prestation de service informatique"}

{f"Détail : {description}" if description else ""}

Durée : {jours} jour(s)
Date : {date_debut.strftime('%d/%m/%Y')}
{f"Lieu : {lieu_mission}" if lieu_mission else "Lieu : Télétravail ou locaux du Client"}
""".strip())

        # RÉMUNÉRATION
        doc.add_heading("ARTICLE 2 - RÉMUNÉRATION", level=1)
        total = contrat.tjm * jours
        doc.add_paragraph(f"""
TJM : {contrat.tjm:,.2f} € HT
Nombre de jours : {jours}
Total : {total:,.2f} € HT

{"TVA non applicable, article 293 B du CGI." if not self.freelance.get('tva_applicable', False) else ""}

Paiement à réception de facture, sous 30 jours.
""".strip())

        # CONDITIONS ESSENTIELLES
        doc.add_heading("ARTICLE 3 - CONDITIONS", level=1)
        doc.add_paragraph("""
- Le Prestataire exerce en toute indépendance
- Les livrables produits sont propriété du Client après paiement
- Le Prestataire respecte la confidentialité des informations
- Le contrat est soumis au droit français
""".strip())

        # SIGNATURES
        doc.add_heading("SIGNATURES", level=1)

        p = doc.add_paragraph()
        p.add_run(f"Fait le {date.today().strftime('%d/%m/%Y')}")

        doc.add_paragraph()

        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Le Prestataire"
        table.cell(0, 1).text = "Le Client"
        table.cell(1, 0).text = "\n\n\n"
        table.cell(1, 1).text = "\n\n\n"
        table.cell(2, 0).text = self.freelance.get('nom', '')
        table.cell(2, 1).text = client.nom

        # Sauvegarde
        filename = f"{contrat.numero}_mission.docx"
        filepath = self.output_dir / filename
        doc.save(str(filepath))

        return str(filepath)

    def _add_standard_articles(self, doc: Document):
        """Ajoute les articles standards (confidentialité, PI, etc.)."""

        # CONFIDENTIALITÉ
        doc.add_heading("ARTICLE 7 - CONFIDENTIALITÉ", level=1)
        doc.add_paragraph("""
Le Prestataire s'engage à garder strictement confidentielles toutes les informations portées à sa connaissance dans le cadre de la mission.

Cette obligation perdure 2 ans après la fin du contrat.
""".strip())

        # PROPRIÉTÉ INTELLECTUELLE
        doc.add_heading("ARTICLE 8 - PROPRIÉTÉ INTELLECTUELLE", level=1)
        doc.add_paragraph("""
Les livrables sont la propriété exclusive du Client dès paiement intégral.

Le Prestataire cède l'ensemble des droits patrimoniaux sur les créations réalisées.
""".strip())

        # LOI APPLICABLE
        doc.add_heading("ARTICLE 9 - LOI APPLICABLE", level=1)
        doc.add_paragraph(f"""
Le présent contrat est soumis au droit français.

Tout litige sera soumis aux tribunaux de {self.freelance.get('ville', 'Strasbourg')}.
""".strip())
